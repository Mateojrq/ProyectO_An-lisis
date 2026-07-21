"""
Genera un documento Word en formato IEEE Conference Paper (DOBLE COLUMNA)
basado en el contenido del proyecto "Análisis de la Educación Superior en Ecuador"

Estructura de secciones Word:
  Section 1: Título, autores, abstract → 1 columna
  Section 2: Cuerpo del artículo → 2 columnas
  (Se insertan section breaks para figuras/tablas anchas que abarcan ambas columnas)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree
import os, copy

IMG_DIR = r'C:\Users\SUPER USUARIO\Downloads\Analisis de datos\repo\docs\images_extracted'
OUTPUT  = r'C:\Users\SUPER USUARIO\Downloads\Analisis de datos\repo\docs\Proyecto_IEEE.docx'

# ──────────────────────────────────────────────
# Utilidades XML para columnas y secciones
# ──────────────────────────────────────────────

def set_columns(sectPr, num_cols, spacing_cm=0.63):
    """Configure number of columns on a w:sectPr element."""
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = parse_xml(f'<w:cols {nsdecls("w")}/>')
        sectPr.append(cols)
    cols.set(qn('w:num'), str(num_cols))
    cols.set(qn('w:space'), str(int(Cm(spacing_cm))))

def insert_continuous_section_break(doc, num_cols, spacing_cm=0.63):
    """Insert a continuous section break into the document and configure columns."""
    # Create a sectPr as a paragraph property (inline section break)
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    sectPr = parse_xml(f'<w:sectPr {nsdecls("w")}><w:type {nsdecls("w")} w:val="continuous"/></w:sectPr>')
    set_columns(sectPr, num_cols, spacing_cm)
    # Copy page size and margins from previous section
    prev_sect = doc.sections[-1]
    pgSz = parse_xml(
        f'<w:pgSz {nsdecls("w")} w:w="{prev_sect.page_width}" w:h="{prev_sect.page_height}"/>'
    )
    sectPr.insert(0, pgSz)
    pgMar = parse_xml(
        f'<w:pgMar {nsdecls("w")} '
        f'w:top="{prev_sect.top_margin}" w:bottom="{prev_sect.bottom_margin}" '
        f'w:left="{prev_sect.left_margin}" w:right="{prev_sect.right_margin}" '
        f'w:header="720" w:footer="720" w:gutter="0"/>'
    )
    sectPr.insert(1, pgMar)
    pPr.append(sectPr)
    # Make the paragraph invisible
    p.text = ''
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    fmt = p.paragraph_format
    fmt.line_spacing = Pt(1)
    return p


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_narrow_margins(section):
    section.top_margin  = Cm(1.91)
    section.bottom_margin = Cm(1.91)
    section.left_margin  = Cm(1.78)   # IEEE ~0.7in
    section.right_margin = Cm(1.78)


# ──────────────────────────────────────────────
# Estilos IEEE
# ──────────────────────────────────────────────

def add_ieee_styles(doc):
    styles = doc.styles

    # --- IEEE Title ---
    s = styles.add_style('IEEE_Title', 1)
    s.font.name = 'Times New Roman'
    s.font.size = Pt(24)
    s.font.bold = True
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.space_after = Pt(12)

    # --- IEEE Author ---
    s = styles.add_style('IEEE_Author', 1)
    s.font.name = 'Times New Roman'
    s.font.size = Pt(11)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.space_after = Pt(2)

    # --- IEEE Abstract label ---
    s = styles.add_style('IEEE_AbstractLabel', 1)
    s.font.name = 'Times New Roman'
    s.font.size = Pt(9)
    s.font.bold = True
    s.font.italic = True
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_before = Pt(12)
    s.paragraph_format.space_after = Pt(6)

    # --- IEEE Abstract text ---
    s = styles.add_style('IEEE_Abstract', 1)
    s.font.name = 'Times New Roman'
    s.font.size = Pt(9)
    s.font.bold = True
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.space_after = Pt(6)
    s.paragraph_format.first_line_indent = Cm(0.5)

    # --- IEEE Section Heading ---
    s = styles.add_style('IEEE_Section', 1)
    s.font.name = 'Times New Roman'
    s.font.size = Pt(10)
    s.font.bold = True
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_before = Pt(12)
    s.paragraph_format.space_after = Pt(6)

    # --- IEEE Subsection Heading ---
    s = styles.add_style('IEEE_Subsection', 1)
    s.font.name = 'Times New Roman'
    s.font.size = Pt(10)
    s.font.bold = True
    s.font.italic = True
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    s.paragraph_format.space_before = Pt(8)
    s.paragraph_format.space_after = Pt(4)

    # --- IEEE Body ---
    s = styles.add_style('IEEE_Body', 1)
    s.font.name = 'Times New Roman'
    s.font.size = Pt(10)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.space_after = Pt(3)
    s.paragraph_format.first_line_indent = Cm(0.35)
    s.paragraph_format.line_spacing = 1.0

    # --- IEEE Figure Caption ---
    s = styles.add_style('IEEE_Caption', 1)
    s.font.name = 'Times New Roman'
    s.font.size = Pt(8)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_before = Pt(4)
    s.paragraph_format.space_after = Pt(8)

    # --- IEEE Reference ---
    s = styles.add_style('IEEE_Reference', 1)
    s.font.name = 'Times New Roman'
    s.font.size = Pt(8)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.space_after = Pt(2)
    s.paragraph_format.left_indent = Cm(0.5)
    s.paragraph_format.first_line_indent = Cm(-0.5)

    # --- IEEE Keywords ---
    s = styles.add_style('IEEE_Keywords', 1)
    s.font.name = 'Times New Roman'
    s.font.size = Pt(9)
    s.font.bold = True
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    s.paragraph_format.space_before = Pt(2)
    s.paragraph_format.space_after = Pt(8)
    s.paragraph_format.first_line_indent = Cm(0.5)

    # --- IEEE Table caption ---
    s = styles.add_style('IEEE_TableCaption', 1)
    s.font.name = 'Times New Roman'
    s.font.size = Pt(8)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_before = Pt(8)
    s.paragraph_format.space_after = Pt(4)


# ──────────────────────────────────────────────
# Funciones de contenido
# ──────────────────────────────────────────────

def add_figure(doc, img_filename, caption, fig_num, width=None):
    """Agrega figura centrada con caption IEEE.
    Para figuras anchas (que abarcan ambas columnas), use add_wide_figure."""
    img_path = os.path.join(IMG_DIR, img_filename)
    if not os.path.exists(img_path):
        doc.add_paragraph(f'[Imagen no disponible: {img_filename}]', style='IEEE_Caption')
        return
    if os.path.getsize(img_path) < 100:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    w = width if width else Inches(3.2)
    try:
        run.add_picture(img_path, width=w)
    except Exception:
        p.text = f'[Error imagen: {img_filename}]'

    cap = doc.add_paragraph(style='IEEE_Caption')
    r1 = cap.add_run(f'Fig. {fig_num}. ')
    r1.bold = True; r1.font.name = 'Times New Roman'; r1.font.size = Pt(8)
    r2 = cap.add_run(caption)
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(8)


def add_wide_figure(doc, img_filename, caption, fig_num, width=None):
    """Inserta una figura que abarca ambas columnas (section break → 1 col → img → section break → 2 col)."""
    insert_continuous_section_break(doc, 1)  # switch to 1 column

    img_path = os.path.join(IMG_DIR, img_filename)
    if not os.path.exists(img_path) or os.path.getsize(img_path) < 100:
        doc.add_paragraph(f'[Imagen no disponible: {img_filename}]', style='IEEE_Caption')
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        w = width if width else Inches(6.5)
        try:
            run.add_picture(img_path, width=w)
        except Exception:
            p.text = f'[Error imagen: {img_filename}]'

    cap = doc.add_paragraph(style='IEEE_Caption')
    r1 = cap.add_run(f'Fig. {fig_num}. ')
    r1.bold = True; r1.font.name = 'Times New Roman'; r1.font.size = Pt(8)
    r2 = cap.add_run(caption)
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(8)

    insert_continuous_section_break(doc, 2)  # back to 2 columns


def add_wide_table(doc, headers, rows, caption, table_num):
    """Tabla que abarca ambas columnas."""
    insert_continuous_section_break(doc, 1)

    cap = doc.add_paragraph(style='IEEE_TableCaption')
    r1 = cap.add_run(f'TABLA {table_num}\n')
    r1.bold = True; r1.font.name = 'Times New Roman'; r1.font.size = Pt(8)
    r2 = cap.add_run(caption.upper())
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(8); r2.font.small_caps = True

    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True; run.font.name = 'Times New Roman'; run.font.size = Pt(7)
        set_cell_shading(cell, "D9E2F3")

    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            if j < n_cols:
                cell = table.rows[i+1].cells[j]; cell.text = ''
                p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(str(val))
                run.font.name = 'Times New Roman'; run.font.size = Pt(7)

    doc.add_paragraph()
    insert_continuous_section_break(doc, 2)


def add_ieee_table(doc, headers, rows, caption, table_num):
    """Tabla dentro de una sola columna."""
    cap = doc.add_paragraph(style='IEEE_TableCaption')
    r1 = cap.add_run(f'TABLA {table_num}\n')
    r1.bold = True; r1.font.name = 'Times New Roman'; r1.font.size = Pt(8)
    r2 = cap.add_run(caption.upper())
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(8); r2.font.small_caps = True

    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]; cell.text = ''
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True; run.font.name = 'Times New Roman'; run.font.size = Pt(7)
        set_cell_shading(cell, "D9E2F3")

    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            if j < n_cols:
                cell = table.rows[i+1].cells[j]; cell.text = ''
                p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(str(val))
                run.font.name = 'Times New Roman'; run.font.size = Pt(7)
    doc.add_paragraph()


def to_roman(n):
    vals = [(1000,'M'),(900,'CM'),(500,'D'),(400,'CD'),
            (100,'C'),(90,'XC'),(50,'L'),(40,'XL'),
            (10,'X'),(9,'IX'),(5,'V'),(4,'IV'),(1,'I')]
    result = ''
    for val, rom in vals:
        while n >= val:
            result += rom
            n -= val
    return result


# ══════════════════════════════════════════════
# CONSTRUCCIÓN DEL DOCUMENTO
# ══════════════════════════════════════════════

def build_document():
    doc = Document()
    add_ieee_styles(doc)

    # ── Sección 1: Página de título (1 columna) ──
    section = doc.sections[0]
    section.page_width  = Cm(21.59)   # Letter
    section.page_height = Cm(27.94)
    set_narrow_margins(section)
    set_columns(section._sectPr, 1)   # explícitamente 1 columna

    # ================================================================
    # TITLE
    # ================================================================
    p = doc.add_paragraph(style='IEEE_Title')
    run = p.add_run('Análisis de la Educación Superior en Ecuador:\nCobertura, Equidad de Acceso y Eficiencia Académica\n(2015-2023)')
    run.font.name = 'Times New Roman'; run.font.size = Pt(24); run.bold = True

    # ================================================================
    # AUTHORS
    # ================================================================
    for text, size, italic in [
        ('Mateo Josue Rodriguez Quevedo', 11, False),
        ('Esteban Alejandro Rodriguez Bazurto', 11, False),
        ('Escuela de Formación de Tecnólogos', 10, True),
        ('Escuela Politécnica Nacional', 10, True),
        ('Quito, Ecuador', 10, True),
        ('Docente: Ing. Juan Pablo Zaldumbide — Período 2026-A', 9, False),
    ]:
        p = doc.add_paragraph(style='IEEE_Author')
        run = p.add_run(text)
        run.font.name = 'Times New Roman'; run.font.size = Pt(size); run.italic = italic

    # ================================================================
    # ABSTRACT (todavía en 1 columna, centrado en toda la página)
    # ================================================================
    abs_label = doc.add_paragraph(style='IEEE_AbstractLabel')
    run = abs_label.add_run('Abstract')
    run.font.name = 'Times New Roman'

    abstract_text = (
        "Ecuador invierte recursos significativos en becas, planta docente e infraestructura de educación superior, "
        "pero no existe una herramienta integrada que relacione la matrícula universitaria con la disponibilidad docente, "
        "la producción científica, la asignación de becas y el flujo de estudiantes proveniente del sistema escolar. "
        "Este proyecto integra siete fuentes de datos abiertos del Ecuador (SENESCYT y Ministerio de Educación) "
        "correspondientes al período 2015-2023: matrícula de educación superior, oferta académica de las IES, "
        "planta docente, artículos publicados en revistas indexadas, becarios, y matrícula/promoción escolar. "
        "Mediante un flujo ETL reproducible en KNIME se estandarizan y unen las fuentes por año, institución y carrera; "
        "en Python se realiza limpieza avanzada, análisis exploratorio y segmentación; y en Power BI se construye un modelo "
        "de datos tipo constelación con medidas DAX que permiten diagnosticar cobertura, equidad de acceso (género, etnia, "
        "discapacidad, provincia) y eficiencia académica. Los resultados buscan aportar evidencia para la toma de decisiones "
        "de política educativa por parte de SENESCYT, el Ministerio de Educación y las propias instituciones de educación superior."
    )
    doc.add_paragraph(abstract_text, style='IEEE_Abstract')

    # Keywords
    kw = doc.add_paragraph(style='IEEE_Keywords')
    r1 = kw.add_run('Palabras clave— ')
    r1.bold = True; r1.italic = True; r1.font.name = 'Times New Roman'; r1.font.size = Pt(9)
    r2 = kw.add_run('educación superior, Ecuador, cobertura, equidad, KNIME, Power BI, análisis de datos, ETL, constelación de hechos.')
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(9); r2.italic = True

    # ── SECTION BREAK → 2 columnas para el cuerpo ──
    insert_continuous_section_break(doc, 2)

    # ================================================================
    # Contadores
    # ================================================================
    sec_num = 0
    sub_counters = {}
    fig_num = 0
    table_num = 0

    def section_heading(title):
        nonlocal sec_num
        sec_num += 1; sub_counters[sec_num] = 0
        p = doc.add_paragraph(style='IEEE_Section')
        run = p.add_run(f'{to_roman(sec_num)}. {title.upper()}')
        run.font.name = 'Times New Roman'

    def subsection_heading(title):
        nonlocal sec_num
        sub_counters[sec_num] = sub_counters.get(sec_num, 0) + 1
        letter = chr(64 + sub_counters[sec_num])
        p = doc.add_paragraph(style='IEEE_Subsection')
        run = p.add_run(f'{letter}. {title}')
        run.font.name = 'Times New Roman'

    def body(text):
        doc.add_paragraph(text, style='IEEE_Body')

    def bullet(text):
        p = doc.add_paragraph(style='IEEE_Body')
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.4)
        run = p.add_run(f'• {text}')
        run.font.name = 'Times New Roman'; run.font.size = Pt(10)

    # Column width for inline figures (~3.3 in inside a column)
    COL_IMG = Inches(3.2)
    COL_IMG_SM = Inches(2.8)

    # ================================================================
    # I. INTRODUCCIÓN
    # ================================================================
    section_heading('Introducción')

    subsection_heading('Contexto y Motivación')
    body(
        'La educación superior es uno de los principales motores de movilidad social y desarrollo '
        'productivo del país. Sin embargo, las decisiones de política pública (apertura de carreras, '
        'asignación de becas, distribución de docentes) suelen tomarse con datos dispersos entre '
        'distintas entidades (SENESCYT, Ministerio de Educación), sin un análisis conjunto que muestre '
        'desbalances entre oferta académica, equidad de acceso y resultados. El Ecuador cuenta con datos '
        'abiertos oficiales y actualizados sobre matrícula, docentes, producción científica, becas y '
        'educación escolar, lo que hace posible construir un diagnóstico integral y con contexto 100% nacional.'
    )

    subsection_heading('Objetivo General')
    body(
        'Diseñar, implementar y defender una solución integral de análisis de datos que combine Python, '
        'KNIME y Power BI para diagnosticar la cobertura, la equidad de acceso y la eficiencia académica '
        'de la educación superior en Ecuador durante el período 2015-2023, integrando información de matrícula, '
        'planta docente, producción científica, becas y educación escolar.'
    )

    subsection_heading('Objetivos Específicos')
    bullet('Extraer y consolidar 7 fuentes de datos abiertos y reales sobre educación superior y escolar en Ecuador.')
    bullet('Diseñar un proceso ETL reproducible en KNIME que estandarice los nombres de instituciones y una las fuentes por año, IES y carrera.')
    bullet('Aplicar limpieza avanzada, análisis exploratorio de datos (EDA) y una técnica de segmentación o clustering en Python.')
    bullet('Construir un modelo de datos tipo constelación en Power BI con medidas DAX de cobertura, equidad y eficiencia académica.')
    bullet('Generar conclusiones y recomendaciones accionables para SENESCYT, el Ministerio de Educación y las IES.')

    subsection_heading('Alcance del Proyecto')
    body(
        'El proyecto cubre el período 2015-2023 (con variaciones según disponibilidad de cada fuente: '
        'docentes hasta 2022, becarios con corte a agosto 2024) a nivel nacional, con desagregación por '
        'provincia, institución y carrera. Incluye educación superior de tercer nivel y su relación con el '
        'sistema escolar previo (bachillerato). No incluye análisis de posgrados internacionales fuera del '
        'programa de becas, ni proyecciones económicas del mercado laboral.'
    )

    # ================================================================
    # II. CASO DE ESTUDIO
    # ================================================================
    section_heading('Caso de Estudio')

    subsection_heading('Descripción del Problema')
    body(
        'Ecuador invierte recursos significativos en becas, planta docente e infraestructura de educación '
        'superior, pero no existe una herramienta integrada que permita ver, en un solo lugar, cómo se '
        'relaciona la matrícula universitaria con la disponibilidad de docentes, la producción científica '
        'de cada institución, la asignación de becas y el flujo de estudiantes que viene desde el sistema '
        'escolar (abandono/promoción en colegios). Las instituciones y el Estado toman decisiones de política '
        'educativa con datos dispersos en distintas fuentes (SENESCYT, MINEDUC), sin un análisis conjunto que '
        'muestre desbalances entre oferta académica, equidad (género, etnia, discapacidad, provincia) y '
        'resultados (publicaciones, becarios, abandono escolar).'
    )

    subsection_heading('Justificación')
    body(
        'El proyecto responde a una necesidad real de política pública: identificar qué instituciones y '
        'provincias tienen menor cobertura o equidad de acceso, qué carreras concentran la oferta académica '
        'sin necesariamente tener más producción científica, y cómo el abandono/promoción en el bachillerato '
        'afecta la matrícula universitaria posterior. Los resultados son útiles para el Ministerio de Educación, '
        'la SENESCYT y las propias IES en la toma de decisiones sobre becas, apertura de carreras y asignación docente.'
    )

    subsection_heading('Preguntas de Análisis')
    body('A continuación se presentan las preguntas guía que orientaron el modelado de datos, la creación de medidas DAX y la construcción del dashboard interactivo:')

    questions = [
        ('¿Cómo ha evolucionado la matrícula de educación superior por año y por tipo de financiamiento (pública/privada) entre 2015 y 2023?',
         'La matrícula en instituciones públicas representa el ~72% de la demanda nacional, manteniendo un crecimiento constante hasta bordear un estancamiento en el periodo post-pandemia.'),
        ('¿Qué provincias presentan menor cobertura de educación superior en relación a su matrícula escolar previa?',
         'Las provincias de la Amazonía (Orellana, Sucumbíos, Pastaza, Zamora Chinchipe) y provincias con alta ruralidad en la Costa (Esmeraldas, Bolívar) muestran la brecha más crítica.'),
        ('¿Cuál es la distribución de la matrícula por sexo, etnia y discapacidad?',
         'Existe paridad de género con ligera mayoría femenina (~52%). La representación indígena y afroecuatoriana es inferior al 7%. Los estudiantes con discapacidad representan menos del 1.8%.'),
        ('¿Qué carreras concentran la mayor demanda estudiantil?',
         'Medicina, Derecho, Administración de Empresas, Psicología e Ingeniería Civil concentran más del 55% de la matrícula.'),
        ('¿Cuál es el ratio estudiantes/docente por institución?',
         'Las IES públicas registran entre 22 y 30 estudiantes por docente, mientras que las IES privadas mantienen promedios de 12 a 18.'),
        ('¿Qué instituciones tienen mayor producción científica por docente?',
         'Las Escuelas Politécnicas (EPN, ESPOL, USFQ, PUCE) lideran, superando 0.6 a 1.2 artículos indexados anuales por docente a tiempo completo.'),
        ('¿Existe correlación entre becarios y producción científica?',
         'Los datos reflejan una correlación positiva moderada (r ≈ 0.65).'),
        ('¿Qué áreas de estudio concentran más becas?',
         'Salud y Bienestar e Ingenierías concentran el mayor volumen. Áreas prioritarias como Ciencias Agropecuarias reciben un porcentaje menor.'),
        ('¿Cómo se relaciona la tasa de abandono escolar con la matrícula posterior?',
         'Existe una relación inversamente proporcional directa entre abandono escolar y matrícula universitaria por provincia.'),
        ('¿Qué diferencias existen entre la oferta académica vigente y la matrícula real?',
         'Se detecta una ineficiencia estructural: carreras "Vigentes" operan con menos del 35% de sus cupos cubiertos.'),
    ]

    for i, (q, a) in enumerate(questions, 1):
        p = doc.add_paragraph(style='IEEE_Body')
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.25)
        r1 = p.add_run(f'{i}) {q} ')
        r1.bold = True; r1.font.name = 'Times New Roman'; r1.font.size = Pt(9)
        r2 = p.add_run(a)
        r2.font.name = 'Times New Roman'; r2.font.size = Pt(9)

    # ================================================================
    # III. ARQUITECTURA DE LA SOLUCIÓN
    # ================================================================
    section_heading('Arquitectura de la Solución')

    subsection_heading('Diagrama de Arquitectura General')
    fig_num += 1
    add_wide_figure(doc, 'image8.png', 'Diagrama de arquitectura general del proyecto.', fig_num, Inches(6.0))

    subsection_heading('Capa de Fuentes de Datos')
    body(
        'Las 7 fuentes provienen del portal de Datos Abiertos del Ecuador (SENESCYT y Ministerio de Educación), '
        'en formatos estructurados: XLSX, CSV y PDF. No se usan APIs ni web scraping en la versión base del proyecto.'
    )

    subsection_heading('Capa de Procesamiento (KNIME y Python)')
    body(
        'KNIME se usa como flujo ETL principal: lectura de las 7 fuentes, estandarización de nombres de instituciones '
        '(mayúsculas, sin tildes/espacios extra), y unión de tablas por las llaves AÑO + NOMBRE_IES (+ NOMBRE_CARRERA '
        'cuando aplica). Python se usa como soporte analítico: limpieza avanzada adicional, análisis exploratorio de '
        'datos y una técnica de segmentación/clustering sobre instituciones o provincias.'
    )

    subsection_heading('Capa de Almacenamiento')
    body('El proyecto utiliza cuatro motores de bases de datos para cumplir con los requisitos del curso:')

    p = doc.add_paragraph(style='IEEE_Body'); p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run('Bases de datos SQL:'); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    bullet('MySQL: almacena docentes_uep, oferta_academica_ies y matricula_uep.')
    bullet('SQLite: almacena registro_matricula_mineduc y descomposicion_matricula_mineduc.')

    p = doc.add_paragraph(style='IEEE_Body'); p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run('Bases de datos NoSQL:'); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    bullet('MongoDB: almacena encuesta_seguimiento_graduados y articulos_publicados.')
    bullet('TinyDB: almacena becarios_agosto_2024_db y becarios_actualizacion_db.')

    subsection_heading('Capa de Visualización')
    body(
        'Power BI se conecta al modelo de datos final (modelo de constelación) y construye un dashboard de 4 páginas: '
        'resumen ejecutivo, exploración, análisis detallado y conclusiones/recomendaciones, con segmentadores de año, '
        'provincia, tipo de financiamiento y perfil demográfico.'
    )

    # ================================================================
    # IV. FUENTES DE DATOS
    # ================================================================
    section_heading('Fuentes de Datos')

    subsection_heading('Descripción y Origen')
    body(
        'Las 7 fuentes provienen del Catálogo de Datos Abiertos del Ecuador (grupo Educación), publicadas por '
        'SENESCYT y el Ministerio de Educación. Todas son reales y verificables públicamente. Fueron obtenidas por '
        'descarga directa desde https://www.datosabiertos.gob.ec/.'
    )

    # Table: Data sources (wide)
    table_num += 1
    add_wide_table(doc,
        ['#', 'Fuente', 'Institución', 'Formato', 'Real/Compl.'],
        [
            ['1', 'Base estadística matrícula UEP 2015-2023', 'SENESCYT', 'XLSX/PDF', 'Real'],
            ['2', 'Oferta académica de las IES', 'SENESCYT', 'XLSX/PDF', 'Real'],
            ['3', 'Base estadística docentes UEP 2015-2022', 'SENESCYT', 'XLSX/PDF', 'Real'],
            ['4', 'Artículos en revistas indexadas', 'SENESCYT', 'XLS/PDF', 'Real'],
            ['5', 'Base de datos de becarios (ago. 2024)', 'SENESCYT', 'CSV/PDF', 'Real'],
            ['6', 'Registro de Matrícula', 'MINEDUC', 'CSV/XLSX', 'Real'],
            ['7', 'Descomposición de la Matrícula', 'MINEDUC', 'CSV/XLSX', 'Real'],
            ['8', 'Becarios - actualización', 'SENESCYT', 'CSV/XLSX', 'Real'],
            ['9', 'Encuesta Seguimiento Graduados', 'SENESCYT', 'CSV', 'Generada'],
        ],
        'Fuentes de datos utilizadas en el proyecto', table_num)

    subsection_heading('Volumen de Registros')
    table_num += 1
    add_wide_table(doc,
        ['Fuente', 'Filas', 'Columnas', 'Nulos detectados'],
        [
            ['Matrícula UEP 2015-2023', '977,964', '19', '0'],
            ['Oferta académica IES', '20,045', '10', '2 (fila basura)'],
            ['Docentes UEP 2015-2022', '38,447', '13', '2'],
            ['Artículos indexados', '87,073', '11', '0'],
            ['Becarios agosto 2024', '65,534', '36', '58'],
            ['Becarios - actualización', '53,325', '36', '~117 (varias cols.)'],
            ['Descomp. Matrícula MINEDUC', '286,111', '23', '0'],
            ['Reg. Matrícula MINEDUC', '290,310', '104', '133,252 (3 cols.)'],
            ['Encuesta Graduados (gen.)', '350,000', '11', '0'],
            ['TOTAL', '2,168,809', '—', '—'],
        ],
        'Volumen de registros por fuente', table_num)

    fig_num += 1
    add_wide_figure(doc, 'image7.png', 'Inspección de datos: vista general de las fuentes cargadas en el entorno Python.', fig_num, Inches(5.5))

    fig_num += 1
    add_wide_figure(doc, 'image13.png', 'Resumen de la inspección de datos por fuente (nulos, columnas, volumen).', fig_num, Inches(5.5))

    # ================================================================
    # V. PROCESO ETL (KNIME)
    # ================================================================
    section_heading('Proceso ETL (KNIME)')

    subsection_heading('Descripción del Workflow')
    body(
        'El flujo de trabajo desarrollado en KNIME implementa un proceso ETL (Extracción, Transformación y Carga) '
        'que integra las 9 fuentes de datos provenientes de las 4 bases de datos creadas (MySQL, SQLite, MongoDB y '
        'TinyDB/CSV), consolidándolas finalmente en un repositorio centralizado en PostgreSQL. El workflow se '
        'organiza en cinco grandes etapas: lectura, limpieza, combinación, transformación y exportación.'
    )

    body(
        '1) Nodos de lectura (extracción): Para cada fuente se utilizó el conector correspondiente '
        'a su motor de origen: MongoDB Connector + MongoDB Reader para colecciones JSON; MySQL Connector + DB '
        'Table Selector + DB Reader para tablas en MySQL; SQLite Connector para tablas del MINEDUC; y CSV Reader '
        'para los archivos de Becarios exportados desde TinyDB.'
    )

    body(
        '2) Nodos de limpieza y preparación: Ungroup para aplanar documentos JSON, GroupBy para agregar información, '
        'Column Renamer para estandarizar nombres, String Manipulation para normalizar texto, Number to String para '
        'homogeneizar tipos, y Duplicate Row Filter para eliminar duplicados.'
    )

    body(
        '3) Nodos de combinación: Concatenate une verticalmente las bases de Becarios (misma estructura), '
        'y Joiner combina horizontalmente las tablas por campos clave comunes.'
    )

    body(
        '4) Nodos de transformación final y exportación: Se aplicaron Column Renamer adicionales. PostgreSQL '
        'Connector + DB Writer escriben el resultado final en PostgreSQL (Neon), generando tres cargas independientes.'
    )

    fig_num += 1
    add_wide_figure(doc, 'image10.png', 'Workflow ETL completo en KNIME Analytics Platform.', fig_num, Inches(6.0))

    fig_num += 1
    add_wide_figure(doc, 'image6.png', 'Detalle de los nodos de carga a PostgreSQL (Neon) y exportación final.', fig_num, Inches(6.0))

    # ================================================================
    # VI. CALIDAD DE DATOS
    # ================================================================
    section_heading('Calidad de Datos')

    subsection_heading('Diagnóstico Inicial')
    body(
        'Antes de iniciar el proceso de limpieza, se realizó un diagnóstico exploratorio de las 9 fuentes de datos '
        'con el fin de identificar valores nulos, duplicados, inconsistencias de formato y posibles outliers.'
    )
    body(
        'Codificación de caracteres: varios archivos CSV presentaron problemas de encoding (tildes y caracteres '
        'especiales mal interpretados), lo que obligó a detectar automáticamente la codificación real de cada archivo '
        '(UTF-8-SIG) antes de poder cargarlos a las bases de datos.'
    )
    body(
        'Nombres de columnas no estandarizados entre fuentes relacionadas: por ejemplo, la base de becarios de '
        'agosto 2024 usa IDENTIFICADOR_BECARIO en mayúsculas, mientras que la actualización usa ID_becario en '
        'minúsculas, requiriendo un mapeo de columnas antes de la unión.'
    )
    body(
        'Alta cardinalidad de columnas: el archivo Registro de Matrícula MINEDUC contiene 104 columnas, muchas '
        'correspondientes a categorías étnicas o de nacionalidad con nombres poco homogéneos.'
    )
    body(
        'Volumen de nulos concentrado: en el Registro de Matrícula MINEDUC, cerca del 46% de los registros no '
        'cuentan con información en las columnas de "estudiantes en edad adecuada al nivel".'
    )

    # ================================================================
    # VII. ANÁLISIS CON PYTHON
    # ================================================================
    section_heading('Análisis con Python')

    subsection_heading('Análisis Exploratorio de Datos (EDA)')
    body(
        'El análisis exploratorio sobre la tabla consolidada fact_educacion_superior (2,771 registros) reveló '
        'los siguientes hallazgos:'
    )

    body(
        'Tendencia temporal: La matrícula mostró un comportamiento irregular entre 2015 y 2016 '
        '(caída de 7.85 a 7.26 millones), seguido de una recuperación sostenida entre 2017 y 2020, '
        'y un salto marcado a partir de 2021 (>9.4 millones). '
        'Este quiebre es consistente con el contexto de pandemia y el aumento de modalidades a distancia.'
    )

    fig_num += 1
    add_figure(doc, 'image17.png', 'Tendencia temporal de la matrícula (2015-2023).', fig_num, COL_IMG)

    body(
        'Distribución por financiamiento: Las instituciones públicas concentran la mayor proporción '
        '(~3.3 millones), seguidas de las cofinanciadas (~2.9 millones), y las autofinanciadas (~500 mil).'
    )

    fig_num += 1
    add_figure(doc, 'image15.png', 'Distribución de la matrícula por financiamiento.', fig_num, COL_IMG)

    body(
        'Concentración institucional: El top 10 está liderado por la PUCE, seguida de la U. de Guayaquil y la '
        'UTPL, evidenciando alta concentración en un grupo reducido de IES.'
    )

    fig_num += 1
    add_figure(doc, 'image18.png', 'Top 10 IES por matrícula.', fig_num, COL_IMG)

    body(
        'Correlaciones: La matriz muestra correlación moderada entre TOTAL_ESTUDIANTES y TOTAL_DOCENTES (0.51), '
        'y entre TOTAL_DOCENTES y TOTAL_ARTICULOS (0.60). La relación matrícula-producción científica es más débil (0.34).'
    )

    fig_num += 1
    add_figure(doc, 'image16.png', 'Matriz de correlación entre variables principales.', fig_num, COL_IMG)

    body(
        'Outliers: El boxplot de TOTAL_ESTUDIANTES evidencia valores atípicos (>40,000) correspondientes a '
        'instituciones grandes con oferta multi-modalidad; no son errores sino concentración real.'
    )

    fig_num += 1
    add_figure(doc, 'image14.png', 'Boxplot de TOTAL_ESTUDIANTES.', fig_num, COL_IMG_SM)

    subsection_heading('Limpieza Avanzada')
    bullet('Tratamiento de outliers mediante IQR para el clustering (sin eliminar del dataset general).')
    bullet('Estandarización de texto en NOMBRE_IES_LIMPIO (espacios, mayúsculas).')
    bullet('Conversión de AÑO a tipo entero para evitar errores en series temporales.')

    # ================================================================
    # VIII. MODELO DE DATOS Y POWER BI
    # ================================================================
    section_heading('Modelo de Datos y Power BI')

    subsection_heading('Modelo de Datos (Constelación)')
    fig_num += 1
    add_wide_figure(doc, 'image2.png', 'Modelo de datos tipo constelación implementado en Power BI.', fig_num, Inches(6.0))

    subsection_heading('Medidas DAX')
    fig_num += 1
    add_figure(doc, 'image4.png', 'Medidas DAX (parte 1).', fig_num, COL_IMG)
    fig_num += 1
    add_figure(doc, 'image5.png', 'Medidas DAX (parte 2).', fig_num, COL_IMG)
    fig_num += 1
    add_figure(doc, 'image3.png', 'Medidas DAX (parte 3).', fig_num, COL_IMG)
    fig_num += 1
    add_figure(doc, 'image11.png', 'Medidas DAX (parte 4).', fig_num, COL_IMG)

    subsection_heading('Diseño del Dashboard')
    fig_num += 1
    add_wide_figure(doc, 'image9.png', 'Vista general del dashboard de Power BI.', fig_num, Inches(6.0))

    subsection_heading('Segmentadores y KPIs')
    body(
        'El dashboard incorpora dos segmentadores (slicers) como filtros globales: Año (2015–2026) y Provincia. '
        'Se implementaron cuatro tarjetas KPI: Total_Matricula_Superior, Total_Becarios, Total_Matricula_Escolar '
        'y Var_%_Matricula_YoY (variación porcentual interanual).'
    )

    # ================================================================
    # IX. RESULTADOS Y VISUALIZACIONES
    # ================================================================
    section_heading('Resultados y Visualizaciones')

    subsection_heading('Resumen Ejecutivo')
    body(
        'La página de Resumen Ejecutivo integra en una sola interfaz la información de Educación Superior '
        '(SENESCYT), el subsistema escolar (MINEDUC) y la ejecución del programa de becas. Se evidencia un '
        'descalce estructural entre graduados de bachillerato y cupos de educación superior, con más del 50% '
        'de matrícula universitaria y becas concentradas en Pichincha y Guayas.'
    )

    subsection_heading('Exploración')
    body(
        'Se observa una acumulación de carreras "No Vigente", sobrecarga en IES públicas tradicionales, '
        'y embudos socioeconómicos y étnicos en la transición hacia la universidad (participación étnica '
        'cae de 14% en secundaria a menos del 7% en educación superior).'
    )

    subsection_heading('Análisis Detallado')
    body(
        'El dashboard integra 10 visualizaciones clave: tarjetas KPI, gráfico de líneas temporal, mapa coroplético '
        'provincial, gráfico de barras por financiamiento, gráfico de anillo por campo de conocimiento, gráfico de '
        'dispersión (producción científica vs. planta docente) y matriz dinámica de becas.'
    )

    subsection_heading('Conclusiones del Dashboard')
    body(
        'Se observa marcada disparidad territorial; el gráfico de dispersión demuestra relación exponencial positiva '
        'entre proporción de profesores con posgrado y producción de artículos indexados. Un volumen considerable de '
        'carreras "No Vigente" genera ineficiencia administrativa.'
    )

    # ================================================================
    # X. CONCLUSIONES
    # ================================================================
    section_heading('Conclusiones')

    body(
        'Centralización Territorial: La oferta académica, planta docente avanzada y recursos de becas se '
        'hiper-concentran en Pichincha y Guayas. Las provincias amazónicas y rurales presentan tasas mínimas '
        'de absorción universitaria.'
    )
    body(
        'Cuello de Botella en la Transición Secundaria-Tercer Nivel: Una fracción sustancial de graduados de '
        'secundaria no logra ingresar al sistema universitario, generando una brecha estructural de capital humano.'
    )
    body(
        'Correlación entre Calidad Docente y Producción Científica: Existe correlación positiva directa entre '
        'docentes titulares acreditados y producción de artículos en revistas indexadas.'
    )
    body(
        'Heterogeneidad y Necesidad de Gobernanza de Datos: La integración de más de 2.1 millones de registros '
        'reveló problemas de estandarización, datos nulos e inconsistencias en las fuentes estatales.'
    )

    # ================================================================
    # XI. RECOMENDACIONES
    # ================================================================
    section_heading('Recomendaciones')

    p = doc.add_paragraph(style='IEEE_Body'); p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run('Para la SENESCYT:'); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    bullet('Rediseñar la fórmula de asignación de becas con ponderador socio-territorial.')
    bullet('Fomentar carreras técnicas cortas (2 años) orientadas a software, datos y agroindustria.')

    p = doc.add_paragraph(style='IEEE_Body'); p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run('Para el MINEDUC:'); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    bullet('Establecer programas de articulación vocacional y nivelación en Bachillerato.')

    p = doc.add_paragraph(style='IEEE_Body'); p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run('Para las IES:'); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    bullet('Automatizar procesos de rediseño curricular ante el CES.')
    bullet('Asignar horas de investigación en la carga horaria docente.')

    # ================================================================
    # XII. PROBLEMAS Y LECCIONES APRENDIDAS
    # ================================================================
    section_heading('Problemas Encontrados y Lecciones Aprendidas')

    p = doc.add_paragraph(style='IEEE_Body'); p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run('Fechas Seriales de Excel: '); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    r2 = p.add_run('Las columnas temporales de becarios contenían números seriales de Excel. Se crearon expresiones '
                    'DAX y transformaciones Python para convertirlos al formato YYYY-MM-DD.')
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(10)

    p = doc.add_paragraph(style='IEEE_Body'); p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run('Codificación UTF-8: '); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    r2 = p.add_run('Archivos con caracteres corruptos. Se aplicaron scripts Python y Power Query para corregirlos.')
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(10)

    p = doc.add_paragraph(style='IEEE_Body'); p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run('Desalineación de Granularidades: '); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    r2 = p.add_run('Se resolvió con arquitectura en Constelación de Hechos usando dimensiones compartidas '
                    '(Dim_Periodo, Dim_Geografia, Dim_IES).')
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(10)

    p = doc.add_paragraph(style='IEEE_Body'); p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run('Lección aprendida: '); r.bold = True; r.italic = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    r2 = p.add_run('El éxito de las visualizaciones y la precisión de los cálculos DAX dependen directamente de '
                    'un diseño sólido del modelo de datos, en lugar de una tabla plana monolítica.')
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(10)

    # ================================================================
    # XIII. REFERENCIAS
    # ================================================================
    section_heading('Referencias')

    references = [
        '[1] SENESCYT, "Base estadística de matrícula UEP 2015-2023," Portal de Datos Abiertos del Ecuador. [En línea]. Disponible: https://www.datosabiertos.gob.ec/. [Accedido: 19-jul-2026].',
        '[2] SENESCYT, "Base de datos de oferta académica de las IES," Portal de Datos Abiertos del Ecuador. [En línea]. Disponible: https://www.datosabiertos.gob.ec/. [Accedido: 19-jul-2026].',
        '[3] SENESCYT, "Base estadística de docentes UEP 2015-2022," Portal de Datos Abiertos del Ecuador. [En línea]. Disponible: https://www.datosabiertos.gob.ec/. [Accedido: 19-jul-2026].',
        '[4] SENESCYT, "Base de datos de artículos publicados en revistas indexadas," Portal de Datos Abiertos del Ecuador. [En línea]. Disponible: https://www.datosabiertos.gob.ec/. [Accedido: 19-jul-2026].',
        '[5] SENESCYT, "Base de datos de becarios, corte agosto 2024," Portal de Datos Abiertos del Ecuador. [En línea]. Disponible: https://www.datosabiertos.gob.ec/. [Accedido: 19-jul-2026].',
        '[6] Ministerio de Educación, "Registro de Matrícula," Portal de Datos Abiertos del Ecuador. [En línea]. Disponible: https://www.datosabiertos.gob.ec/. [Accedido: 19-jul-2026].',
        '[7] Ministerio de Educación, "Descomposición de la Matrícula," Portal de Datos Abiertos del Ecuador. [En línea]. Disponible: https://www.datosabiertos.gob.ec/. [Accedido: 19-jul-2026].',
    ]

    for ref in references:
        doc.add_paragraph(ref, style='IEEE_Reference')

    # ================================================================
    # ANEXOS
    # ================================================================
    doc.add_page_break()
    p = doc.add_paragraph(style='IEEE_Section')
    run = p.add_run('ANEXOS'); run.font.name = 'Times New Roman'

    p = doc.add_paragraph(style='IEEE_Subsection')
    run = p.add_run('A. Repositorio en GitHub'); run.font.name = 'Times New Roman'
    body('El proyecto técnico completo se encuentra en: https://github.com/Mateojrq/ProyectO_An-lisis.git')

    p = doc.add_paragraph(style='IEEE_Subsection')
    run = p.add_run('B. Videos Explicativos'); run.font.name = 'Times New Roman'
    body('Video 1: Proceso de datos y ETL (4-6 min). Video 2: Análisis del dashboard y conclusiones (4-6 min).')

    # ── Configurar la ÚLTIMA sección del documento también a 2 columnas ──
    last_sect = doc.sections[-1]
    set_columns(last_sect._sectPr, 2)

    # ================================================================
    # GUARDAR
    # ================================================================
    doc.save(OUTPUT)
    print(f"✅ Documento IEEE (doble columna) generado exitosamente en:")
    print(f"   {OUTPUT}")
    print(f"   Total de figuras: {fig_num}")
    print(f"   Total de tablas: {table_num}")


if __name__ == '__main__':
    build_document()
